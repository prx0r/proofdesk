"""Doctavian integration tests — payload builder, branch logic, gate/envelope boundary."""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from src.providers.doctavian import build_generation_payload
from src.models.domain import (
    Case, Document, ExtractedFact, Assertion, Resolution,
    CaseState, AssertionResult, ExceptionSeverity, ResolutionDecision,
)
from src.engine.orchestrator import run_pipeline
from src.state.machine import can_request_signature


def _record(with_fail=True, with_resolution=False):
    assertions = [
        {"predicate": "quote.total == platform + support", "result": "PASS", "detail": "ok", "severity": "BLOCKER", "rule_version": "arith-v1"},
    ]
    if with_fail:
        assertions.append({
            "predicate": "insurance.expiry_date >= required_coverage_until",
            "result": "FAIL", "detail": "31-day gap", "severity": "BLOCKER",
            "rule_version": "coverage-v1",
        })
    resolutions = []
    if with_resolution:
        resolutions.append({"decision": "CONDITIONAL_ACCEPT", "reason": "Renewed cert required", "actor_id": "pm"})
    return {
        "case_id": "case_t1", "record_id": "rec_t1", "content_hash": "sha256:abc123",
        "facts": [
            {"field": "vendor.legal_name", "value_normalized": "Northstar Data Systems Ltd."},
            {"field": "quote.total", "value_normalized": "42500"},
        ],
        "assertions": assertions, "resolutions": resolutions,
    }


def test_payload_builder():
    # CLEARED band when no failures
    p = build_generation_payload(_record(with_fail=False))
    assert p["risk_band"] == "CLEARED", p["risk_band"]
    assert p["has_conditions"] == "false" and p["condition_count"] == 0

    # CONDITIONAL with failure + resolution
    p = build_generation_payload(_record(with_resolution=True))
    assert p["risk_band"] == "CONDITIONAL"
    assert p["condition_count"] == 1
    assert p["failed_checks"][0]["rule"] == "coverage-v1"
    assert p["resolutions"][0]["actor"] == "pm"

    # ESCALATED: failure, no resolution
    p = build_generation_payload(_record())
    assert p["risk_band"] == "ESCALATED"

    # Explicit confidence from foxit module wins
    p = build_generation_payload(_record(), {"confidence": 0.41, "band": "CONDITIONAL",
        "field_risks": [{"field": "insurance.expiry_date", "detail": "6% > 3% budget"}]})
    assert p["signing_confidence"] == "0.41"
    assert p["condition_count"] == 2  # check fail + field risk
    assert any(c["predicate"].startswith("confidence:") for c in p["failed_checks"])

    # Determinism: same input → same payload (except generated_date)
    a = build_generation_payload(_record()); b = build_generation_payload(_record())
    a.pop("generated_date"); b.pop("generated_date")
    assert a == b

    print("[PASS] DOCT-001..005: payload builder bands, confidence wiring, determinism")


def test_template_exists_and_has_branches():
    path = os.path.join(os.path.dirname(__file__), "..", "data", "templates", "vendor_approval_memo.docx")
    assert os.path.exists(path), "template missing"
    from docx import Document
    doc = Document(os.path.abspath(path))
    text = "\n".join(par.text for par in doc.paragraphs)
    for marker in ["{!risk_band}", "mdoc:repeater", "mdoc:paragraph",
                   "{!condition_count}", "{!#fc#.predicate}"]:
        assert marker in text, f"template missing {marker}"
    print("[PASS] DOCT-006: template v2 contains branch/loop markers")


def test_envelope_gate_boundary():
    """Envelope creation is forbidden before gate passes — premature attempt denies."""
    case = Case()
    for name in ["procurement_request.pdf", "vendor_quote.pdf",
                 "insurance_certificate.pdf", "security_questionnaire.pdf"]:
        path = os.path.join(os.path.dirname(__file__), "..", "data", "test_pdfs", name)
        if os.path.exists(path):
            # doc_id aliased to stub extraction pattern ("certificate_insurance")
            stub_id = "certificate_insurance" if name == "insurance_certificate.pdf" else name.replace(".pdf", "")
            with open(path, "rb") as f:
                case.documents.append(Document(
                    doc_id=stub_id, case_id=case.case_id,
                    filename=name, content_type="application/pdf",
                    raw_bytes=f.read()))
    run_pipeline(case)  # lands in REVIEW_REQUIRED (insurance gap)

    gate = can_request_signature(case)
    assert not gate["allowed"], "gate must deny pre-resolution"
    codes = {r["code"] for r in gate["reasons"]}
    assert "UNRESOLVED_BLOCKER" in codes and "INVALID_STATE" in codes
    print(f"[PASS] DOCT-007: envelope blocked pre-gate ({len(gate['reasons'])} reasons)")


if __name__ == "__main__":
    test_payload_builder()
    test_template_exists_and_has_branches()
    test_envelope_gate_boundary()
    print("\nDoctavian integration tests: ALL PASS")
